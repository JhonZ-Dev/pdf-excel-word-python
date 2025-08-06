from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
# === Función para colorear fondo de celda ===
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
# --- 1. Cargar plantilla membretada con una tabla dummy (debe existir en el Word)
doc = Document("plantilla.docx")
def set_cell_bg(cell, color_hex):
    """Set background color of a cell by direct XML injection."""
    cell._tc.get_or_add_tcPr().append(
        parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{}"/>'.format(color_hex))
    )

# --- 2. Borra la tabla dummy de la plantilla, para que NO aparezca en el archivo final ---
if doc.tables:
    tbl = doc.tables[0]
    tbl._element.getparent().remove(tbl._element)

# --- 3. Crea la tabla general (7 columnas) y aplica Table Grid ---
tabla = doc.add_table(rows=0, cols=7)
tabla.style = 'Table Grid'

# 1. PROFORMA (cabecera)
row = tabla.add_row().cells
row[0].merge(row[6])
p = row[0].paragraphs[0]
run = p.add_run('PROFORMA 12345')
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Arial' 
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_cell_bg(row[0], 'dae9f7')  # Fondo azul claro

# 2. Campos principales (campo + respuesta)
campos = [
    ('CLIENTE', 'FISCALIA DE BOLIVAR'),
    ('RUC', '0260018030001'),
    ('DIRECCIÓN:', 'Av. Candido Rada y S/N'),
    ('FECHA', '2025-08-03'),
    ('TELEFONO', '022345678'),
    ('NECESIDAD', 'Compra de útiles'),
    ('FUNCIONARIO ENCARGADO:', 'Lic. María Pérez'),
    ('CORREO', 'maría.perez@correo.com'),
    ('HORA MAXIMA:', '12:00')
]
for label, valor in campos:
    row = tabla.add_row().cells
    row[0].merge(row[1])
    set_cell_bg(row[0], 'e8f2fe')
    p = row[0].paragraphs[0]
    p.clear()
    run = p.add_run(label)
    run.font.size = Pt(10)
    run.bold=True
    run.font.name = 'Arial'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    row[2].merge(row[6])
    row[2].text = valor

# 3. OBJETO DE COMPRA (título y respuesta)
row = tabla.add_row().cells
row[0].merge(row[6])
row[0].text = "OBJETO DE COMPRA"
row[0].paragraphs[0].runs[0].bold = True

row = tabla.add_row().cells
row[0].merge(row[6])
row[0].text = "ADQUISICIÓN SUMINISTROS DE OFICINA PARA STOCK BODEGA DE LA FISCALIA DE BOLIVAR"

# 4. CABECERAS DE ÍTEMS
headers = ['No.', 'CPC', 'UNIDAD', 'ESPECIFICACIONES', 'CANTIDAD', 'P. UNIT', 'P.TOTAL']
row = tabla.add_row().cells
for i, h in enumerate(headers):
    row[i].text = h
    row[i].paragraphs[0].runs[0].bold = True

# 5. FILAS DE ÍTEMS (ejemplo)
items = [
    (1, '12345', 'CAJA', 'Lápiz HB', 10, 2.5, 25),
    (2, '54321', 'UNIDAD', 'Borrador', 20, 0.5, 10),
]
for item in items:
    row = tabla.add_row().cells
    for i, value in enumerate(item):
        row[i].text = str(value)

# 6. TOTAL
row = tabla.add_row().cells
row[0].merge(row[5])
row[0].text = "TOTAL"
row[6].text = "$35.00"

# 7. NO GRAVAMOS IVA
row = tabla.add_row().cells
row[0].merge(row[6])
row[0].text = "NO GRAVAMOS IVA - SOMOS REGIMEN RIMPE - NEGOCIO POPULAR"

# 8. CAMPOS FINALES
finales = [
    ('PLAZO DE ENTREGA', '7 días'),
    ('VIGENCIA DE LA OFERTA', '15 días'),
    ('FORMA DE PAGO:', 'Transferencia'),
    ('METODOLOGIA DE TRABAJO', 'Entrega a bodega'),
    ('GARANTIA', '1 año'),
    ('ENLACE', 'https://ejemplo.com')
]
for label, valor in finales:
    row = tabla.add_row().cells
    row[0].merge(row[1])
    row[0].text = label
    row[2].merge(row[6])
    row[2].text = valor

# 9. Firma fuera de la tabla
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DAYANA LISBETH ZAMBRANO MACIAS\nREPRESENTANTE LEGAL\nRUC: 2350621211001')
run.bold = True
run.font.size = Pt(14)

doc.save('proforma_final.docx')
